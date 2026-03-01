#!/usr/bin/env python3
"""
Erstellt ein Word-Dokument Template für docsautomator
mit der Roots Stückliste Formatierung
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_table_border(table):
    """Fügt Rahmen zu einer Tabelle hinzu"""
    tbl = table._element
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '666666')
        tblBorders.append(border)

    tblPr.append(tblBorders)

def create_template():
    doc = Document()

    # Titel
    title = doc.add_heading('ROOTS SYSTEMKONFIGURATOR', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.color.rgb = RGBColor(46, 160, 67)  # Grün
    title_run.font.size = Pt(24)
    title_run.font.bold = True

    subtitle = doc.add_heading('Stückliste', level=2)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.color.rgb = RGBColor(46, 160, 67)

    doc.add_paragraph()

    # Projekt-Information Box
    info_table = doc.add_table(rows=1, cols=1)
    info_table.style = 'Light Grid Accent 1'
    cell = info_table.rows[0].cells[0]
    cell.text = ''

    # Füge Projekt-Info hinzu
    p = cell.add_paragraph()
    run = p.add_run('Projekt: ')
    run.bold = True
    p.add_run('{{Projektname}}')

    p = cell.add_paragraph()
    run = p.add_run('Exportdatum: ')
    run.bold = True
    p.add_run('{{Exportdatum}}')

    p = cell.add_paragraph()
    p.add_run()  # Leerzeile

    p = cell.add_paragraph()
    run = p.add_run('GEBÄUDEINFORMATIONEN')
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(46, 160, 67)

    p = cell.add_paragraph()
    run = p.add_run('Name: ')
    run.bold = True
    p.add_run('{{Gebaeude_Name}}')

    p = cell.add_paragraph()
    run = p.add_run('Baujahr: ')
    run.bold = True
    p.add_run('{{Gebaeude_Baujahr}}')

    p = cell.add_paragraph()
    run = p.add_run('Adresse: ')
    run.bold = True
    p.add_run('{{Gebaeude_Strasse}} {{Gebaeude_Hausnummer}}')

    p = cell.add_paragraph()
    run = p.add_run('Stockwerke: ')
    run.bold = True
    p.add_run('{{Gebaeude_Stockwerke}}')

    doc.add_paragraph()
    doc.add_paragraph()

    # Komponenten Überschrift
    heading = doc.add_heading('KOMPONENTEN', level=1)
    heading_run = heading.runs[0]
    heading_run.font.color.rgb = RGBColor(46, 160, 67)
    heading_run.font.size = Pt(16)

    # Komponenten Tabelle
    komp_table = doc.add_table(rows=2, cols=8)
    komp_table.style = 'Light Grid Accent 1'
    add_table_border(komp_table)

    # Header Row
    header_cells = komp_table.rows[0].cells
    headers = ['Pos.', 'Name', 'Modultyp', 'Hersteller', 'Abmessungen', 'Menge', 'Preis', 'Gesamt (€)']
    for i, header_text in enumerate(headers):
        cell = header_cells[i]
        cell.text = header_text
        # Formatiere Header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(46, 160, 67)
        # Hintergrundfarbe
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), 'E8F5E9')
        cell._element.get_or_add_tcPr().append(shading_elm)

    # Data Row mit docsautomator Platzhaltern (Line Items Syntax!)
    data_cells = komp_table.rows[1].cells
    placeholders = [
        '{{line_items_1}}{{Position}}',  # Line Items beginnt im ersten Feld
        '{{Name}}',
        '{{Modultyp}}',
        '{{Hersteller}}',
        '{{Abmessungen}}',
        '{{Menge}} {{Einheit}}',
        '{{Preis_pro_Einheit}}€',
        '{{Gesamtpreis}}€'
    ]
    for i, placeholder in enumerate(placeholders):
        data_cells[i].text = placeholder
        for paragraph in data_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10)

    # Zwischensumme Komponenten
    p = doc.add_paragraph()
    run = p.add_run('Zwischensumme Komponenten: ')
    run.bold = True
    run.font.size = Pt(12)
    run2 = p.add_run('{{Komponenten_Summe}}€')
    run2.bold = True
    run2.font.size = Pt(12)
    run2.font.color.rgb = RGBColor(46, 160, 67)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph()
    doc.add_paragraph()

    # Leitungen Überschrift
    heading = doc.add_heading('LEITUNGEN (ROHRE & KABEL)', level=1)
    heading_run = heading.runs[0]
    heading_run.font.color.rgb = RGBColor(46, 160, 67)
    heading_run.font.size = Pt(16)

    # Leitungen Tabelle
    leit_table = doc.add_table(rows=2, cols=8)
    leit_table.style = 'Light Grid Accent 1'
    add_table_border(leit_table)

    # Header Row
    header_cells = leit_table.rows[0].cells
    headers = ['Pos.', 'Von', 'Zu', 'Typ', 'Länge', 'Dimension', 'Preis/m (€)', 'Gesamt (€)']
    for i, header_text in enumerate(headers):
        cell = header_cells[i]
        cell.text = header_text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(46, 160, 67)
        # Hintergrundfarbe
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), 'E8F5E9')
        cell._element.get_or_add_tcPr().append(shading_elm)

    # Data Row mit docsautomator Platzhaltern (Line Items Syntax!)
    data_cells = leit_table.rows[1].cells
    placeholders = [
        '{{line_items_2}}{{Position}}',  # Line Items 2 für zweite Liste
        '{{Von_Modul}} ({{Von_Ausgang}})',
        '{{Zu_Modul}} ({{Zu_Eingang}})',
        '{{Verbindungstyp}}',
        '{{Laenge_m}} m',
        '{{Dimension}}',
        '{{Preis_pro_m}}€',
        '{{Gesamtpreis}}€'
    ]
    for i, placeholder in enumerate(placeholders):
        data_cells[i].text = placeholder
        for paragraph in data_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10)

    # Zwischensumme Leitungen
    p = doc.add_paragraph()
    run = p.add_run('Zwischensumme Leitungen: ')
    run.bold = True
    run.font.size = Pt(12)
    run2 = p.add_run('{{Leitungen_Summe}}€')
    run2.bold = True
    run2.font.size = Pt(12)
    run2.font.color.rgb = RGBColor(46, 160, 67)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph()
    doc.add_paragraph()

    # Gesamtsumme Box
    summe_table = doc.add_table(rows=1, cols=1)
    cell = summe_table.rows[0].cells[0]

    # Hintergrundfarbe für Gesamtsumme
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), '2EA043')
    cell._element.get_or_add_tcPr().append(shading_elm)

    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('GESAMTSUMME: ')
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(255, 255, 255)
    run2 = p.add_run('{{Gesamtsumme}}€')
    run2.bold = True
    run2.font.size = Pt(20)
    run2.font.color.rgb = RGBColor(255, 255, 255)

    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Komponenten: {{Komponenten_Summe}}€  |  Leitungen: {{Leitungen_Summe}}€')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(255, 255, 255)

    # Speichern
    doc.save('/Users/florianhackl-kohlweiss/roots-configurator/Roots_Stueckliste_Template.docx')
    print('✅ Template erfolgreich erstellt: Roots_Stueckliste_Template.docx')

if __name__ == '__main__':
    create_template()
